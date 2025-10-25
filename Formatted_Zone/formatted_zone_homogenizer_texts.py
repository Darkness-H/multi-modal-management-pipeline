# Importing useful dependencies
import io
import os
import time
import chardet
import docx
import fitz
import zipfile
from bs4 import BeautifulSoup
from odf.opendocument import load
from odf.text import P
from odf.teletype import extractText as odf_text
from striprtf.striprtf import rtf_to_text
import re
import logging
import markdown

logger = logging.getLogger(__name__)



def _norm(s: str) -> str:
    """
    Normalize paragraph spacing in a text block.

    This helper function standardizes newline characters, removes leading/trailing
    whitespace, and ensures that paragraphs are separated by exactly one blank line.

    s          : str                   - Input text to normalize.
    """
    if not s:
        return ""
    s = s.replace("\r\n", "\n").replace("\r", "\n").strip()
    # Collapse 3 or more newlines into 2 (blank line between paragraphs)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s


def extract_from_docx(body: bytes) -> str:
    """
    Extract plain text from a DOCX document, preserving paragraph breaks.

    The function reads a DOCX file from raw bytes, collects non-empty paragraph
    texts in order, and joins them with a blank line between paragraphs.
    Final spacing is normalized via `_norm()`.

    This extracts `document.paragraphs` only.
    Text inside tables, headers/footers, text boxes/shapes, and comments is not included.

    body          : bytes                   - The raw bytes of a `.docx` file.
    """
    doc = docx.Document(io.BytesIO(body))

    # Gather non-empty paragraphs in order, trimming incidental whitespace.
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    return _norm("\n\n".join(paragraphs))

def extract_from_pdf(body: bytes) -> str:
    """
    Extract plain text from a PDF by aggregating page text blocks.

    This approach uses PyMuPDF's block-level extraction to produce a more
    reliable paragraph grouping than line-by-line text.

    body          : bytes                   - The raw bytes of a `.pdf` file.
    """
    doc = fitz.open(stream=io.BytesIO(body), filetype="pdf")
    pages = []

    try:
        for i in range(len(doc)):

            # Each block is a tuple; index 0..4 commonly: (x0, y0, x1, y1, text, ...)
            blocks = doc.load_page(i).get_text("blocks")

            # Sort blocks by top (y0) then left (x0) to approximate natural reading order
            blocks.sort(key=lambda b: (b[1], b[0]))

            # Collect non-empty block texts; strip incidental whitespace
            txts = [b[4].strip() for b in blocks if b[4].strip()]
            # Separate blocks on the page with one blank line
            pages.append("\n\n".join(txts))

    finally:
        doc.close()


    return _norm("\n\n".join(pages))



def extract_from_zip(body: bytes) -> str:
    """
    Extract plain text from an EPUB or ZIP archive containing HTML/XHTML files.

    EPUB files are essentially ZIP archives with structured HTML content.
    This function scans the archive for all `.xhtml`, `.html`, or `.htm` files,
    extracts their readable text, and joins them in lexical order (sorted by
    filename) to approximate the book's original reading order.

    body          : bytes                   - Raw bytes of an EPUB or ZIP archive.

    """

    parts = []

    with zipfile.ZipFile(io.BytesIO(body)) as z:
        for name in sorted(z.namelist()):
            # Only process HTML/XHTML content files
            if name.lower().endswith((".xhtml", ".html", ".htm")):
                with z.open(name) as f:
                    # Parse HTML and extract visible text
                    soup = BeautifulSoup(f.read(), "html.parser")
                    text = soup.get_text("\n").strip()
                    if text:
                        parts.append(text)

    return _norm("\n\n".join(parts))


def extract_from_html(body: bytes) -> str:
    """
    Extract plain text from an HTML document while preserving structural line breaks.

    The function parses HTML content, strips away all markup, and retrieves only
    the visible text. A newline (`\\n`) is inserted between HTML elements to
    roughly preserve the visual separation of blocks (e.g., paragraphs, divs, headings).
    The resulting text is then normalized using `_norm()` to ensure clean paragraph spacing.

    body          : bytes                   -  Raw bytes of the HTML file or response content.
    """

    soup = BeautifulSoup(body, "html.parser")

    # Extract visible text from all tags, separating each block with a newline
    text = soup.get_text(separator="\n")

    return _norm(text)


def extract_from_odt(body: bytes) -> str:
    """
    Extract plain text from an OpenDocument Text (.odt) file, preserving paragraph breaks.

    This function opens an ODT document from raw bytes, iterates over all
    `<text:p>` paragraph elements, extracts their visible text using
    `teletype.extractText`, and joins non-empty paragraphs with a blank line.
    The final text is normalized via `_norm()` to ensure consistent spacing.

    body          : bytes                   -  Raw bytes of the `.odt` file.

    """

    # Open the ODT document directly from bytes
    doc = load(io.BytesIO(body))

    # Collect non-empty paragraph texts in order
    paragraphs = []
    for p in doc.getElementsByType(P):
        text = odf_text(p).strip()
        if text:
            paragraphs.append(text)

    return _norm("\n\n".join(paragraphs))


# RTF
def extract_from_rtf(body: bytes) -> str:
    """
    Extract plain text from an RTF document.

    The function tries to use `striprtf` (if installed) for robust RTF parsing.
    If `striprtf` is unavailable or parsing fails, it falls back to a simple
    heuristic that:
      - Replaces paragraph markers (e.g., ``\\par``) with blank lines,
      - Removes control words/groups and braces,
      - Leaves best-effort plain text.

    body          : bytes                   -  Raw bytes of the `.rtf` file.
    """
    try:
        # Preferred path: use striprtf if available
        text = rtf_to_text(body.decode("utf-8", errors="ignore"))

    # Fall through to heuristic cleanup
    except Exception:
        # Heuristic fallback: very rough RTF cleanup
        s = body.decode("utf-8", errors="ignore")

        # Replace paragraph control words with double newlines
        # \par and \pard → paragraph breaks
        s = re.sub(r"\\par[d]?", "\n\n", s)

        # Remove control groups like {\*\...} or {\fonttbl ...}, keep a space
        s = re.sub(r"{\\[^}]*}", " ", s)

        # Remove standalone control words (e.g., \b, \i, \fs24, \cf1, \u1234)
        s = re.sub(r"\\[a-zA-Z]+-?\d*", " ", s)

        # Remove remaining braces
        s = re.sub(r"[{}]", " ", s)
        text = s

    return _norm(text)


def extract_from_md(body: bytes) -> str:
    """
    Extract plain text from Markdown content by converting to HTML first.

    Converts Markdown to HTML using the standard 'markdown' module,
    then removes markup with BeautifulSoup, preserving paragraph spacing.

    body          : bytes                   -  Raw bytes of the `.md` file.
    """
    html = markdown.markdown(body.decode("utf-8", errors="ignore"))
    soup = BeautifulSoup(html, "html.parser")
    text = soup.get_text(separator="\n")
    return _norm(text)

# map each extension to its extractor
EXT_MAP = {
    ".docx": extract_from_docx,
    ".pdf":  extract_from_pdf,
    ".epub": extract_from_zip,    # epub is a ZIP of XHTML/HTML
    ".zip":  extract_from_zip,
    ".html": extract_from_html,
    ".odt":  extract_from_odt,
    ".rtf":  extract_from_rtf,
    ".md":   extract_from_md,
}


def convert_and_replace(key, body):
    """
    Convert a supported document (PDF, DOCX, HTML, etc.) to plain text (.txt).

    This function performs content conversion only:
      - It detects the file extension,
      - Selects the correct extractor function from EXT_MAP,
      - Extracts plain text from the raw file bytes,
      - Returns both the converted text and the new filename.

    key             : str                   - Object key (path) within the bucket.
    body            : bytes                 - Raw file bytes to be converted.
    """

    # Determine the file extension and corresponding extractor function
    ext = os.path.splitext(key)[1].lower()
    extractor = EXT_MAP.get(ext)
    if extractor is None:
        logger.warning("Skipping unsupported file: %s", key)
        return False, None, None

    try:
        content = extractor(body)
        name, _ = os.path.splitext(key)
        new_key = name + ".txt"
        logger.debug("Converted %s -> %s", key, new_key)
        return True, content, new_key
    except Exception as e:
        logger.exception("Failed to extract %s: %s", key, e)
        return False, None, None


def convert_documents_to_txt(client, bucket, prefix=""):
    """
    Scan all objects under a given S3 prefix, normalize text encodings to UTF-8,
    and convert non-text documents (PDF, DOCX, HTML, etc.) into plain text (.txt).


    client          : obj                   - S3-compatible client (e.g., boto3.client("s3")).
    bucket          : str                   - Target S3/MinIO bucket.
    prefix          : str                   - Optional key prefix (acts like a folder path).

    """
    paginator = client.get_paginator("list_objects_v2")

    t0 = time.perf_counter()
    summary = {
        "scanned": 0,
        "converted_txt": 0,   # txt files re-encoded to utf-8
        "converted_other": 0, # non-txt docs converted to .txt
        "skipped": 0,         # skipped unsupported or already UTF-8 or empty file
        "failed": 0,
    }
    pages = paginator.paginate(Bucket=bucket, Prefix=prefix)

    for page in pages:
        for obj in page.get("Contents", []):

            if obj['Size'] == 0:  # Skip the folder itself (if the file size is 0)
                summary["skipped"] += 1
                continue

            key = obj["Key"]
            summary["scanned"] += 1

            try:
                # Get the file object from S3
                resp = client.get_object(Bucket=bucket, Key=key)
                body = resp["Body"].read()  # Read the file content

                # Case 1: plain text file (check encoding)
                if key.endswith(".txt"):
                    # Use chardet to detect the file encoding
                    result = chardet.detect(body)
                    current_encoding = result['encoding']

                    # Skip if the file is already in UTF-8 encoding
                    if current_encoding in ("utf-8", "ascii", None):
                        logger.debug("Already UTF-8 or ASCII: %s", key)
                        summary["skipped"] += 1
                        continue

                    logger.debug("Converting %s from %s to UTF-8", key, current_encoding)

                    # Decode the content using the detected encoding and re-encode it in UTF-8
                    content = body.decode(current_encoding, errors='ignore')  # Ignore characters that can't be decoded

                    # Upload the converted content back to S3 in UTF-8
                    client.put_object(
                        Bucket=bucket,
                        Key=key,  # Make sure the file key (path) is correct
                        Body=content.encode('utf-8'),
                        ContentType="text/plain"
                    )
                    summary["converted_txt"] += 1
                    logger.debug("Re-encoded %s successfully", key)

                else:
                    handled, content, new_key = convert_and_replace(key, body)
                    if handled:
                        client.put_object(
                            Bucket=bucket,
                            Key=new_key,
                            Body=content.encode("utf-8"),
                            ContentType="text/plain",
                        )
                        summary["converted_other"] += 1
                        logger.debug("Converted %s → %s", key, new_key)

                    else:
                        summary["skipped"] += 1
                        logger.debug("Skipped unsupported extension: %s", key)

            except Exception as e:
                summary["failed"] += 1
                logger.exception("Failed to process %s: %s", key, e)

    elapsed= time.perf_counter() - t0
    logger.info(
        "Texts conversion completed in %.2fs — scanned=%d, converted_txt=%d, converted_others=%d, skipped=%d, failed=%d",
        elapsed, summary["scanned"], summary["converted_txt"], summary["converted_other"], summary["skipped"], summary["failed"]
    )
